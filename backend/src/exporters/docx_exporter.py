from io import BytesIO

from docx import Document

from models.resume_model import Resume


class DOCXExporter:
    """Exports a resume to a DOCX file in memory."""

    def export(self, resume: Resume) -> bytes:
        buffer = BytesIO()
        document = Document()

        document.add_heading(resume.title, level=1)

        if resume.target_role:
            document.add_heading("Cargo Alvo", level=2)
            document.add_paragraph(resume.target_role)

        if resume.content:
            document.add_heading("Conteúdo do Currículo", level=2)
            document.add_paragraph(resume.content)

        document.add_heading("Versão", level=2)
        document.add_paragraph(str(resume.version))

        document.save(buffer)

        docx_bytes = buffer.getvalue()
        buffer.close()

        return docx_bytes
